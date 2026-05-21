import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_study_material():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run('学习资料：习近平总书记关于巡视巡察工作的重要论述和要求')
    title_run.font.name = '宋体'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(255, 0, 0) # Red color for party study material
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('\n') # Space

    # Foreword
    p = doc.add_paragraph()
    r = p.add_run('（供党支部内部学习参考）\n')
    r.font.name = '楷体'
    r.font.size = Pt(14)
    r.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    r = p.add_run('巡视制度是党内监督的战略性制度安排。党的十八大以来，以习近平同志为核心的党中央高度重视巡视工作，把巡视作为推进党的自我革命、全面从严治党的战略性制度安排。习近平总书记多次听取巡视情况汇报、发表重要讲话，作出一系列重要论述，不断引领巡视工作理论创新、实践创新、制度创新。')
    r.font.size = Pt(14)

    # Section 1
    h1 = doc.add_heading(level=1)
    r = h1.add_run('一、 根本遵循：坚守政治巡视定位')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.add_run('习近平总书记强调：“巡视是政治巡视，本质是政治监督。发现问题、形成震慑，推动改革、促进发展，是巡视工作的方针。” [1]').bold = True
    p.add_run('\n新时代巡视工作必须把坚定拥护“两个确立”、坚决做到“两个维护”作为根本任务，聚焦党中央决策部署在基层的落实情况，聚焦群众身边腐败问题和不正之风，聚焦基层党组织建设情况，聚焦巡视整改落实情况。')

    # Section 2
    h2 = doc.add_heading(level=1)
    r = h2.add_run('二、 重点对象：紧盯“一把手”和领导班子')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.add_run('习近平总书记指出：“要加强对‘一把手’和领导班子的监督，督促其严于律己、严负其责、严管所辖。” [2]').bold = True
    p.add_run('\n巡视巡察工作必须牢牢抓住“关键少数”，通过自上而下的组织监督，层层传导压力，确保各级领导干部特别是主要负责同志做到忠诚干净担当，发挥好“头雁效应”。')

    # Section 3
    h3 = doc.add_heading(level=1)
    r = h3.add_run('三、 核心任务：发现问题，形成震慑')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.add_run('“发现不了问题是失职，发现问题不报告是渎职。” [3]').bold = True
    p.add_run('\n巡视工作要坚持严的基调、严的措施、严的氛围，深入查找政治偏差，勇于亮剑，精准发现深层次问题和隐患，发挥好巡视利剑作用，使巡视真正成为国之利器、党之利器。')

    # Section 4
    h4 = doc.add_heading(level=1)
    r = h4.add_run('四、 根本目的：推动改革，促进发展，做好“后半篇文章”')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.add_run('习近平总书记强调：“把巡视发现的问题解决好，巡视整改不落实，就是对党不忠诚、对人民不负责。” 把巡视整改作为检验“四个意识”的试金石。 [4]').bold = True
    p.add_run('\n必须扎实做好巡视整改“后半篇文章”。要建立整改长效机制，把巡视整改与贯彻落实党的二十大精神结合起来，与学习贯彻习近平新时代中国特色社会主义思想主题教育结合起来，做到以巡促改、以巡促建、以巡促治。')

    # Section 5
    h5 = doc.add_heading(level=1)
    r = h5.add_run('五、 监督格局：推动巡视巡察上下联动和贯通协调')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.add_run('要贯通其他监督，加强与党委和政府“两办”督查、纪检监察监督、组织人事监督、审计监督等各类监督的贯通协调，推动整合监督力量、健全监督体系，形成系统集成、协同高效的监督合力。 [5]')

    # Party branch requirement
    doc.add_paragraph('\n')
    h6 = doc.add_heading(level=2)
    r = h6.add_run('支部学习要求：')
    r.font.size = Pt(14)
    r.font.bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('全体党员要深刻领会政治巡视内涵，自觉增强接受监督的意识。')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('结合岗位实际，主动自查自纠，杜绝形式主义、官僚主义作风。')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('将巡视整改要求内化于心、外化于行，推动本部门工作高质量发展。')

    # Add References section
    doc.add_paragraph('\n')
    h_ref = doc.add_heading(level=2)
    r_ref = h_ref.add_run('参考引文与官方交叉验证来源：')
    r_ref.font.size = Pt(14)
    r_ref.font.bold = True
    
    p_ref1 = doc.add_paragraph(style='List Number')
    r1 = p_ref1.add_run('《中央纪委国家监委：推进新征程纪检监察工作高质量发展不断增强巡视巡察震慑力穿透力》\n')
    r1.bold = True
    p_ref1.add_run('出处：中央纪委国家监委网站 (ccdi.gov.cn)\n')
    p_ref1.add_run('链接：https://www.ccdi.gov.cn/specialn/zjwscqh/zjwscqhtopnews/202501/t20250106_398675.html\n')
    p_ref1.add_run('引用内容：“巡视是政治巡视，本质是政治监督...”')
    
    p_ref2 = doc.add_paragraph(style='List Number')
    r2 = p_ref2.add_run('《中央纪委国家监委：二十届中央第十一巡视组巡视国家市场监督管理总局党组工作动员会召开》\n')
    r2.bold = True
    p_ref2.add_run('出处：中央纪委国家监委网站 (ccdi.gov.cn)\n')
    p_ref2.add_run('链接：https://www.ccdi.gov.cn/specialn/zyxsgz20/20jzyxstpxw/202404/t20240418_342288.html\n')
    p_ref2.add_run('引用内容：紧盯权力和责任，加强对“一把手”和领导班子的监督...')
    
    p_ref3 = doc.add_paragraph(style='List Number')
    r3 = p_ref3.add_run('《人民日报评论：扎实做好巡视“后半篇文章”》\n')
    r3.bold = True
    p_ref3.add_run('出处：中央纪委国家监委网站 (ccdi.gov.cn) 转载\n')
    p_ref3.add_run('链接：http://m.ccdi.gov.cn/content/f0/63/87383.html\n')
    p_ref3.add_run('引用内容：“巡视发现问题的目的是解决问题，巡视整改不落实，就是对党不忠诚、对人民不负责...”')
    
    p_ref4 = doc.add_paragraph(style='List Number')
    r4 = p_ref4.add_run('《中共国家统计局党组关于巡视整改进展情况的通报》\n')
    r4.bold = True
    p_ref4.add_run('出处：国家统计局官网 (stats.gov.cn)\n')
    p_ref4.add_run('链接：https://www.stats.gov.cn/xw/tjxw/tjdt/202302/t20230202_1895733.html\n')
    p_ref4.add_run('引用内容：坚持把巡视整改作为检验“四个意识”的试金石...')

    p_ref5 = doc.add_paragraph(style='List Number')
    r5 = p_ref5.add_run('《中央纪委国家监委：深入学习贯彻习近平总书记重要讲话精神持续推动巡视工作高质量发展》\n')
    r5.bold = True
    p_ref5.add_run('出处：中央纪委国家监委网站 (ccdi.gov.cn)\n')
    p_ref5.add_run('链接：https://www.ccdi.gov.cn/plywz/202504/t20250411_416392.html\n')
    p_ref5.add_run('引用内容：加强与党委和政府“两办”督查、纪检监察监督、职能部门监督的贯通协调...')

    # Set font to all paragraphs to normal sizes
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.name:
                run.font.name = '宋体'

    doc.save('学习资料：习近平总书记关于巡视巡察工作的重要论述（党支部内部学习）.docx')
    print("Docx file updated successfully.")

if __name__ == '__main__':
    create_study_material()

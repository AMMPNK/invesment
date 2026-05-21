#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

BLUE_DARK  = RGBColor(0x1F, 0x36, 0x64)
BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)
BLUE_EN    = RGBColor(0x1F, 0x6B, 0xB5)
ORANGE_ZH  = RGBColor(0xC0, 0x55, 0x14)
GREY       = RGBColor(0x60, 0x60, 0x60)

def sp(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def h1(text):
    p = doc.add_heading(text, level=1)
    sp(p, before=14, after=4)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    sp(p, before=10, after=4)
    return p

def en(text):
    p = doc.add_paragraph()
    sp(p, before=2, after=2, line=14)
    t = p.add_run("[EN]  ")
    t.bold = True; t.font.color.rgb = BLUE_EN; t.font.size = Pt(10)
    b = p.add_run(text)
    b.font.size = Pt(10)
    return p

def zh(text):
    p = doc.add_paragraph()
    sp(p, before=2, after=7, line=14)
    t = p.add_run("[ZH]  ")
    t.bold = True; t.font.color.rgb = ORANGE_ZH; t.font.size = Pt(10)
    b = p.add_run(text)
    b.font.size = Pt(10)
    return p

def note(text):
    p = doc.add_paragraph()
    sp(p, before=4, after=4)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.italic = True; r.font.color.rgb = GREY
    return p

def blt(text):
    p = doc.add_paragraph(style='List Bullet')
    sp(p, before=1, after=1, line=13)
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p

def hr():
    p = doc.add_paragraph()
    sp(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bot); pPr.append(pBdr)
    return p

def tbl(headers, rows, col_widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        if hcells[i].paragraphs[0].runs:
            r = hcells[i].paragraphs[0].runs[0]
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hcells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E75B6'); tcPr.append(shd)
    for row_data in rows:
        rc = t.add_row().cells
        for i, ct in enumerate(row_data):
            rc[i].text = ct
            if rc[i].paragraphs[0].runs:
                rc[i].paragraphs[0].runs[0].font.size = Pt(10)
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(col_widths[i])
    doc.add_paragraph()

# ===== COVER =====
c1 = doc.add_paragraph()
c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(c1, before=30, after=4)
r = c1.add_run("Chapter 5: Collaborative Principled Negotiation")
r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = BLUE_DARK

c2 = doc.add_paragraph()
c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(c2, before=0, after=4)
r2 = c2.add_run("第五章：合作原则谈判法")
r2.font.size = Pt(16); r2.font.bold = True; r2.font.color.rgb = BLUE_MID

c3 = doc.add_paragraph()
c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(c3, before=4, after=20)
r3 = c3.add_run("白远《国际商务谈判》英文版·第六版 | 中英双语整理 & NotebookLM PPT 大纲\n中国人民大学出版社")
r3.font.size = Pt(10); r3.font.italic = True; r3.font.color.rgb = GREY

hr()

pa = doc.add_paragraph()
sp(pa, before=10, after=6)
ra = pa.add_run("PART A  原文核心内容摘要（中英对照）")
ra.font.size = Pt(14); ra.font.bold = True; ra.font.color.rgb = BLUE_DARK

# ===== 5.1 =====
h1("5.1  Collaborative Principled Negotiation and Its Four Components  合作原则谈判法及其四大组成部分")

en("The increasing acceptance of the win-win concept has brought forth the development of new negotiation theories. A representative one is Collaborative Principled Negotiation (CPN), also commonly known as Harvard Principled Negotiation, established by Roger Fisher and William Ury, professors from Harvard University. Their key work Getting to Yes: Negotiating Agreement Without Giving In is widely regarded as the Bible of negotiations.")
zh("随着双赢理念的广泛认可，新型谈判理论不断涌现。其中最具代表性的是合作原则谈判法（CPN），又称哈佛原则谈判法，由哈佛大学教授罗杰·费希尔与威廉·尤里创立。其代表作《谈判力》（Getting to Yes）被誉为谈判领域的圣经。")

en("Core Definition: The core of CPN is to reach a solution beneficial to both parties by stressing interests and values, NOT by haggling. It decides issues on their merits, looks for mutual gains, and insists that any agreement be based on fair standards independent of the will of either side. CPN is hard on the merits, soft on the people. It employs no tricks and no posturing.")
zh("核心定义：CPN 的精髓在于通过强调各方的利益与价值（而非讨价还价）达成双赢方案。它根据问题是非曲直做出决定，寻求互利，并坚持结果必须基于独立于任何一方意志的公平标准。精髓：对问题强硬，对人温和，不耍手段，不摆姿态。")

ph = doc.add_paragraph()
sp(ph, before=6, after=4)
rh = ph.add_run("四大基本要素 / Four Basic Components:")
rh.bold = True; rh.font.size = Pt(11)

tbl(
    ["#", "Component / 要素", "Principle / 原则"],
    [
        ["1", "People  人", "Separate the people from the problem  把人与问题分开"],
        ["2", "Interests  利益", "Focus on interests, not positions  着眼于利益，而非立场"],
        ["3", "Gaining  方案", "Invent options for mutual gain  创造双赢方案"],
        ["4", "Criteria  标准", "Introduce objective criteria  引入客观评判标准"],
    ],
    [1.0, 4.2, 8.3]
)
note("The four components are interrelated and should be applied to the whole course of the negotiations.\n四要素相互关联，贯穿谈判全程。")
hr()

# ===== 5.2 =====
h1("5.2  Separate the People from the Problem  对事不对人")

en("It is generally understood that negotiations go smoothly in a friendly and sincere atmosphere. Unfortunately, high tension is often built up due to negotiators' prejudice against the other party or misled interpretation of the other party's intentions. Negotiations can degenerate into personal disputes. A basic fact: you are not dealing with abstract representatives, but with human beings who have emotions, deeply held values, different backgrounds and viewpoints -- and they are unpredictable.")
zh("通常认为谈判应在友好真诚的氛围中推进。然而现实中，由于偏见、误解等原因，谈判往往充满紧张对立，甚至演变为人身攻击。基本事实：你面对的不是抽象代表，而是有情感、有价值观、有不同背景的活生生的人，且他们的行为无法预测。")

en("To find your way through people problems, think in terms of three basic categories: Perception, Emotion, and Communication.")
zh("处理人际问题，需从三个基本维度入手：认知（Perception）、情绪（Emotion）、沟通（Communication）。")

h2("5.2.1  Perception  认知")

en("Understanding the other side's thinking is not simply a useful activity -- their thinking IS the problem. Conflict lies not in objective reality, but in people's heads.")
zh("理解对方的思维方式，不仅仅是有益的——他们的思维方式本身就是问题所在。冲突并不存在于客观现实中，而存在于人的头脑中。")

en("Put yourself in their shoes: People tend to pick out facts that confirm their prior perceptions and disregard those that challenge them. The ability to see the situation as the other side sees it is one of the most important skills a negotiator can possess.")
zh("换位思考：人们总倾向于关注能印证自身既有认知的信息，而忽略对其构成挑战的信息。能够从对方角度看待问题，是谈判者最重要的技能之一。")

en("Do not blame them for your problem: Blaming is counterproductive. Under attack, the other side becomes defensive and resists what you say.")
zh("不要因为自己的问题而责备对方：责备适得其反——对方会变得防御甚至反击。")

en("Give the other side a stake in the outcome: The feeling of participation in the process is perhaps the single most important factor in determining whether a negotiator accepts a proposal. In a sense, the process is the product.")
zh("让对方参与到结果中来：参与感是决定谈判者是否接受方案的最重要因素之一。在某种意义上，过程即产品。")

h2("5.2.2  Emotion  情绪")

en("In a bitter dispute, feelings may be more important than talk. Emotions on one side generate emotions on the other. Emotions may quickly bring a negotiation to an impasse.")
zh("在激烈争议中，情绪往往比言辞更重要。一方的情绪会引发另一方的情绪连锁反应，从而迅速将谈判推向僵局。")

en("Allow the other side to let off steam: The best strategy while the other side lets off steam is to listen quietly without responding to their attacks. People obtain psychological release through the simple process of recounting their grievances.")
zh("允许对方发泄：对方发泄时，最佳策略是安静倾听，不回应攻击。通过倾诉委屈，人们可以获得心理释放。")

en("Do not react to emotional outbursts: An effective technique adopted in the 1950s by the Human Relations Committee in the steel industry -- adopt the rule that only one person can get angry at a time. This legitimizes others not responding stormily and helps people control their emotions.")
zh("不要对情绪爆发做出情绪性反应：钢铁行业劳资关系委员会曾采用的有效技巧——规定同一时刻只能有一人发怒。这为他人不做激烈回应提供了正当理由，也有助于双方控制情绪。")

h2("5.2.3  Communication  沟通")

en("Three big problems in communication: (1) Negotiators may not be talking to each other in a way to be understood. (2) Even if talking directly and clearly, they may not be listening. (3) The third problem is misunderstanding.")
zh("沟通的三大障碍：(1)双方可能没有以能被理解的方式真正交流；(2)即便在直接清晰地表达，对方也可能没在倾听；(3)误解。")

en("Listen actively and acknowledge what is being said: The cheapest concession you can make to the other side is to let its negotiators know they have been heard.")
zh("积极倾听并认可对方所说的内容：你能给对方最廉价的让步，就是让对方感到被倾听。")

en("Speak about yourself, not about them: It is more persuasive to describe a problem in terms of its impact on you than in terms of what the negotiator of the other side did or why. Say 'I feel let down' instead of 'You are a racist.'")
zh("谈论自己，而非评价对方：用陈述某件事对你的影响来描述问题更有说服力。说我感到失望，而非你是个种族主义者。")

en("Avoid trying to score points and debating them as opponents: It is unpersuasive to blame the other party, engage in name-calling, or raise your voice.")
zh("避免争辩和给对方打分：指责对方、人身攻击、提高嗓门，都是无说服力的行为。")

note("Summary 5.2: Treat your counterpart as a cooperator sitting on the same boat -- sinking and floating together, moving toward mutual success hand in hand.\n小结：把对方视为同乘一船的合作伙伴——同舟共济，携手迈向共同成功。")
hr()

# ===== 5.3 =====
h1("5.3  Focus on Interests But Not Positions  着眼于利益，而非立场")

en("Position refers to the preset attitude towards something or a choice of sides. Classic story (Fisher & Ury): Two men quarrel in a library about how much to open a window. A librarian asks why: one wants fresh air, the other wants no draft. She opens a window in the next room -- both interests satisfied. This story is typical: since parties' goals are to agree on positions, they naturally reach an impasse.")
zh("立场是对某件事预先持有的态度或选择。经典案例——图书馆开窗之争：两位读者为开窗多少争执不休。图书管理员分别询问原因——一人要新鲜空气，另一人要避免穿堂风——于是打开隔壁房间的窗户，两个利益同时满足。这个故事揭示了谈判的典型困境：当双方都执着于立场时，僵局往往难以避免。")

en("Successful negotiations result from mutual giving and taking of interests, not keeping firm on positions. The method works because: (1) There is always more than one way to fulfill each other's interests. (2) Both sides can always find common interests -- otherwise they would not sit together at all.")
zh("成功的谈判源于双方在利益层面的相互给予与接纳。聚焦共同利益之所以有效：(1)满足双方利益的方式不止一种；(2)双方总能找到某些共同利益——否则根本不会坐到谈判桌前。")

h2("5.3.1  Identify Interests  明确利益")

en("Explore their interests which stood in our way: Ask 'Why?' to understand the needs, hopes, fears, or desires that the position serves. A position is concrete and explicit; interests underlying it may be unexpressed, intangible, and inconsistent.")
zh("探寻阻碍我方的对方利益：不断追问为什么——目的是理解立场背后的需求、希望、恐惧或愿望。立场往往具体而明确，其背后的利益却可能是隐性的、无形的甚至是矛盾的。")

en("Examine the different interests of different people on their side: Almost every negotiation has many interests. A common error is assuming everyone on the other side has the same interests.")
zh("审视对方不同人的不同利益：几乎每场谈判都不止一种利益诉求。一个常见错误是假设对方所有人都有相同的利益。")

en("Look at their human needs underlying their positions: Look for bedrock concerns that motivate all people -- security, economic well-being, a sense of belonging, recognition, and control over one's life. In many negotiations, we tend to think the only interest involved is money.")
zh("透过立场看到背后的人性需求：特别关注驱动所有人的基本需求——安全感、经济利益、归属感、被认可感、以及对自己生活的掌控感。许多时候，我们错误地认为对方的唯一利益就是金钱。")

h2("5.3.2  Talk about Interests  讨论利益")

en("The chance of achieving your interests increases when you communicate them. Make your interests come alive; acknowledge their interests as part of the problem.")
zh("将利益清晰传达出去，才能提高实现利益的机会。让你的利益活起来，将对方的利益作为共同问题的一部分加以认可。")

en("Be hard on the problem, soft on the people: Commit yourself to your interests, not to your position. Often the wisest solutions -- maximum gain for you at minimal cost to the other side -- are produced only by strongly advocating your interests.")
zh("对问题强硬，对人温和：坚守利益，而非立场。最聪明的解决方案往往只有通过强力倡导自身利益才能实现。")

hr()

# ===== 5.4 =====
h1("5.4  Invent Options for Mutual Gain  创造双赢方案")

en("The third component provides an approach to the fulfillment of the two parties' demands. Why are negotiators easily trapped? Many negotiations focus on a single event and the solution is either winning or losing. The distributive nature of interest gaining limits scope of thinking.")
zh("第三个要素为满足双方需求提供了具体路径。为何谈判者容易陷入困境？许多谈判聚焦于单一事件，结果非赢即输。利益分配的零和特性限制了人们的思维空间。")

en("Three major barriers to inventing creative options: (1) Fixed plan -- both sides perceive the size of the cake is fixed: my gain is your loss. (2) Seeking only one solution -- negotiators are inclined to rest on their laurels without exploring other options. (3) Considering only one's own options -- a successful negotiation is a process of giving and taking; options should be a consolidated body of both sides' interests.")
zh("创造性方案生成面临三大障碍：(1)固定馅饼假设——双方都觉得利益总量固定，你多我就少；(2)只寻求单一答案——谈判者急于找到那个答案，不愿意探索其他选项；(3)只考虑满足自身需求的方案——成功的谈判是相互给予的过程，方案必须综合双方利益。")

h2("5.4.1  Invent Creative Options  发明创造性方案")

en("(1) Separate inventing options from evaluating them: Since judgment hinders imagination, separate the creative act from the critical one. Invent first, decide later. Consider a brainstorming session -- the key ground rule is to postpone all criticism and evaluation. With those inhibitions removed, one idea stimulates another, like firecrackers setting off one by one.")
zh("(1) 将创造方案与评价方案分开：评判会束缚想象力，因此要将创意行为与批判行为分开。先发明，后决定。组织头脑风暴会议——核心规则是暂停一切批评与评价。当束缚被解除后，一个想法激发另一个想法，如鞭炮一串串炸响。")

en("(2) Develop several options before looking for a solution: The key to wise decision-making lies in selecting from a great number and variety of options. Sources: find out the problems first, analyze second, consider what to do next, and finally come up with specific feasible suggestions.")
zh("(2) 在寻找解决方案之前先发展多个选项：明智决策的关键在于从大量多样的选项中筛选。路径：先发现问题、分析问题、思考对策，最终提出具体可行的行动建议。")

h2("5.4.2  Look for Mutual Gain  寻求互利")

en("A major block to creative problem-solving lies in the assumption of a fixed pie. Actually, there almost always exists the possibility of joint gain.")
zh("创造性解决问题的最大障碍，是固定馅饼的假设。然而实际上，几乎总存在共同获益的可能。")

en("(1) Identify shared interests: Three points worth remembering: First, shared interests lie latent in every negotiation. Second, shared interests are opportunities, not godsends -- make them explicit as a concrete, future-oriented shared goal. Third, stressing shared interests makes the negotiation smoother. Classic metaphor: Passengers in a lifeboat afloat in the middle of the ocean with limited rations will subordinate their differences over food in pursuit of their shared interests in getting to shore.")
zh("(1) 找到共同利益。三点牢记：第一，共同利益潜伏在每一场谈判之中；第二，共同利益是机遇不是天赐之物——要将其明确化、目标化，使其具体且面向未来；第三，强调共同利益能让谈判更顺畅。经典比喻：同舟共济的乘客，会因上岸这一共同利益而搁置食物分配上的分歧。")

en("(2) Look for options that will make the decision easier for them: Since success depends upon the other side's decision, make that decision as painless as possible for them.")
zh("(2) 寻找能让对方更容易做决定的选项：你的成功有赖于对方的决定，因此要让对方的决定变得尽可能无痛苦。")

hr()

# ===== 5.5 =====
h1("5.5  Introduce Objective Criteria  引入客观评判标准")

en("The first three components advocate considering both parties' interests and designing constructive patterns. However, conflicts and disputes will not disappear no matter how creative the two sides are. When they cannot decide which option is reasonable and rational, looking for an objective criterion will be a way out.")
zh("前三个要素都在倡导兼顾双方利益、设计建设性方案。然而，无论双方多么富有创意，冲突与争议都不会消失。当双方无法判断哪个方案合理、理性时，寻找客观评判标准便成为出路。")

en("Classic Case -- India vs. US at an Ocean Law Conference: India insisted on $60 million initial fee per site; US rejected firmly, suggesting no initial fee. Stalemate ensued. A representative found the MIT model for deep-seabed mining economics -- accepted as an objective criterion by both sides. The model showed India's proposal would interrupt a company's normal operation (fee due 5 years before revenue). India agreed to reconsider; the model also showed US that a reasonable initial fee was economically viable. Result: US gave up its position.")
zh("经典案例——国际海洋法会议上的印美谈判：印度坚持每个开采点收取6000万美元初始费用；美国坚决反对，主张零初始费用。陷入僵局后，一位代表发现了麻省理工学院（MIT）的深海采矿经济模型，被双方接受为客观标准。该模型显示印度提议的初始费用会打断公司正常运营（需在产生收益前5年缴费）。印度同意重新考虑；模型同样告知美方，合理数额的初始费用在经济上是可行的。结果：美方放弃了原来的立场。")

h2("5.5.1  Developing Objective Criteria  建立客观标准")

en("(1) Look for fair standards. Three qualities of an objective criterion: (i) Independent of the wills of all parties and free from sentimental influence. (ii) Legitimate and practical (e.g., a river as boundary rather than three yards east of the riverbank). (iii) At least theoretically accepted by both sides, as with the MIT model. Different issues need different criteria: price negotiations -- cost, market value, depreciation; other negotiations -- expert opinions, international conventions, legal documents.")
zh("(1) 寻找公平标准。客观标准的三个特质：(i)独立于所有各方的意志之外，不受情感因素影响；(ii)正当且可行（如以河流为边界，而非河岸以东三码）；(iii)至少在理论上被双方所接受，MIT模型正是如此。不同议题有不同标准：价格谈判——成本、市场行情、折旧；其他谈判——专家意见、国际惯例、法律文件。")

en("(2) Look for fair procedures: To produce an outcome independent of wills, use fair standards or fair procedures. Example: when one party cuts a cake, ask the other party to choose first. Other fair procedures: doing it in turns or drawing lots. Drawing lots has an inherent fairness -- results may be unequal, but each side has an equal opportunity.")
zh("(2) 寻求公平程序：除公平标准外，还可以使用公平程序来解决利益冲突。例如：一方切蛋糕，另一方先挑选。其他公平程序：轮流、抽签。抽签本身具有内在公平性——结果可能不均等，但每一方都有平等的机会。")

h2("5.5.2  Standards for Successful Negotiations  成功谈判的三项标准")

en("In Fisher and Ury's view, three standards can be applied to judging success or failure of a negotiation approach: (1) An agreement should satisfy the legitimate interests of both parties to the maximum and resolve their conflicts, while protecting public interests. (2) The agreement should be highly efficient. (3) The agreement will improve, or at least not hurt, the relationship of the two parties.")
zh("费希尔与尤里认为，可用以下三项标准判断一种谈判方式的成败：(1)如果能达成协议，应最大限度地满足双方的正当利益、解决冲突，同时维护公众利益；(2)协议应当高效；(3)协议应当改善，或至少不损害双方的关系。")

note("CPN Final Summary: CPN applies to almost all situations -- diplomacy, business, law, personal life. Unlike other strategies, if the other side learns CPN too, it does not become more difficult -- it becomes easier. Its success relies not on tricks, but on fairness, objectiveness, and mutual understanding.\nCPN总结：合作原则谈判法适用于几乎所有情境。与其他策略不同，对方也学会了只会让谈判变得更容易。其成功不依赖于耍手段，而依赖于公平、客观与相互理解。")

# ===== PAGE BREAK + PART B =====
doc.add_page_break()

pb = doc.add_paragraph()
sp(pb, before=10, after=6)
rb = pb.add_run("PART B  PPT 大纲（NotebookLM 生成版）")
rb.font.size = Pt(14); rb.font.bold = True; rb.font.color.rgb = BLUE_DARK

note("使用说明：将本文档上传至 NotebookLM，输入提示词：\nBased on this structured outline, generate a 12-slide presentation with bullet points and speaker notes in both English and Chinese for each slide.")

SLIDES = [
    ("Slide 1", "封面 | Title Slide", [
        "Title: Chapter 5 -- Collaborative Principled Negotiation",
        "Subtitle: How to Win Without Fighting -- A Harvard Approach",
        "Source: Bai Yuan, International Business Negotiation (6th Ed.), RUC Press",
    ]),
    ("Slide 2", "开场案例 | Opening Case: China-US Agricultural Negotiation", [
        "Background: China's WTO accession -- wheat import dispute",
        "China banned US wheat from 7 northwest states due to smut concerns",
        "China's move: lifted the ban -- created political goodwill for WTO accession",
        "Result: Win-win -- market access resolved + WTO support gained",
        "CPN Insight: Focus on interests (WTO entry), not positions (import ban)",
    ]),
    ("Slide 3", "什么是CPN？| What Is CPN?", [
        "Founded by Roger Fisher & William Ury, Harvard University",
        "Key book: Getting to Yes -- The Bible of Negotiations",
        "Core: Reach win-win by stressing interests & values, NOT haggling",
        "Motto: Hard on the merits, soft on the people",
        "Applies to: diplomacy, business, law, personal disputes",
    ]),
    ("Slide 4", "四大核心要素 | The 4 Components Overview", [
        "(1) People -- Separate the people from the problem  把人与问题分开",
        "(2) Interests -- Focus on interests, not positions  着眼于利益，而非立场",
        "(3) Gaining -- Invent options for mutual gain  创造双赢方案",
        "(4) Criteria -- Introduce objective criteria  引入客观评判标准",
        "Key: Interrelated -- apply to the WHOLE course of negotiations",
    ]),
    ("Slide 5", "要素一：对事不对人 | Separate People from the Problem", [
        "Human beings have emotions, values, different viewpoints -- they are unpredictable",
        "3 categories: Perception, Emotion, Communication",
        "Perception: Put yourself in their shoes; do not blame them",
        "Emotion: Let them let off steam; only one angry person at a time",
        "Communication: Listen actively; speak about yourself, not them",
        "Key: Process is the product -- involve the other side in the process",
    ]),
    ("Slide 6", "要素二：着眼利益 | Focus on Interests, Not Positions", [
        "Classic: Library Window Story -- same positions, different underlying interests",
        "Position = what you say you want | Interest = why you want it",
        "Behind opposed positions lie shared and compatible interests",
        "How: Ask Why -- uncover needs, hopes, fears, desires",
        "Human needs: security, belonging, recognition, control",
        "Speak up: Be hard on the problem, soft on the people",
    ]),
    ("Slide 7", "要素三：创造双赢 | Invent Options for Mutual Gain", [
        "3 Barriers: (1) Fixed pie assumption  (2) Seeking only one solution  (3) Considering only own options",
        "Solution 1: Brainstorming -- separate inventing from evaluating",
        "Key rule: Postpone ALL criticism during idea generation",
        "One idea sets off another -- like firecrackers",
        "Solution 2: Expand the pie -- make it larger before cutting it",
        "Find shared interests + look for options that make their decision easier",
    ]),
    ("Slide 8", "要素四：客观标准 | Introduce Objective Criteria", [
        "When creativity is not enough -- we need objective criteria",
        "3 qualities: Independent of wills / Legitimate and practical / Theoretically accepted by both",
        "Sources: market value, expert opinion, MIT-type models, law, industry norms",
        "Case: MIT deep-seabed mining model resolved India-US stalemate",
        "Fair procedures: one cuts other chooses; drawing lots; taking turns",
        "Key: Yield to principle, not to pressure",
    ]),
    ("Slide 9", "成功谈判的三重标准 | 3 Standards for Successful Negotiations", [
        "Fisher & Ury's 3 standards:",
        "(1) Agreement satisfies legitimate interests of both parties + protects public interests",
        "(2) Agreement is highly efficient",
        "(3) Agreement improves or at least does not hurt the parties' relationship",
        "Takeaway: Shifts negotiation from win-lose to value creation + relationship building",
    ]),
    ("Slide 10", "互动讨论 | Case Discussion & Reflection", [
        "Q1: How did the HK/Macao handover negotiations embody CPN's four components?",
        "Q2: When personally attacked in a negotiation, how do you apply Section 5.2?",
        "Q3: Give examples of how joint efforts can make the interest cake bigger.",
        "Q4: What other criteria can serve as fair procedural standards beyond those mentioned?",
    ]),
    ("Slide 11", "总结 | Conclusion -- Hard on the Problem, Soft on the People", [
        "(1) People: Separate -- treat counterpart as cooperator on the same boat",
        "(2) Interests: Focus -- ask Why not just What",
        "(3) Gaining: Invent -- brainstorm, expand the pie, look for mutual gain",
        "(4) Criteria: Insist -- yield to principle, not to pressure",
        "CPN's success relies on fairness, objectiveness, and mutual understanding",
        "If the other side learns this too -- it becomes EASIER, not harder",
    ]),
    ("Slide 12", "课后思考题 | Questions for Discussion and Consideration", [
        "Q1: Disputes over sovereignty (e.g., HK/Macao handover) are often hard to address. Please apply CPN to analyze the key factors for success.",
        "Q2: What criteria in your view will be considered as fair procedural standards in addition to the ones already mentioned?",
        "Q3: Is it easy to separate people from problems? What will be your reaction when you are attacked personally during a negotiation?",
        "Q4: Please give some examples to illustrate that through joint efforts of negotiating parties a bigger cake of interests can be made.",
    ]),
]

for slide_num, slide_title, bullets in SLIDES:
    ph2 = doc.add_paragraph()
    sp(ph2, before=10, after=3)
    rnum = ph2.add_run("[" + slide_num + "]  ")
    rnum.bold = True; rnum.font.size = Pt(12); rnum.font.color.rgb = BLUE_MID
    rtitle = ph2.add_run(slide_title)
    rtitle.bold = True; rtitle.font.size = Pt(12); rtitle.font.color.rgb = BLUE_DARK
    for b in bullets:
        blt(b)

# ===== SAVE =====
out = "/Users/baixiaoyang/Desktop/Chapter5_CPN_双语整理与PPT大纲.docx"
doc.save(out)
print("saved: " + out)

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 OpenClaw 商业化策略研究报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面设置 A4 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 颜色常量 ──
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # 深墨
C_RED    = RGBColor(0xC0, 0x39, 0x2B)   # 强调红
C_GRAY   = RGBColor(0x55, 0x55, 0x66)   # 次要文字
C_LGRAY  = RGBColor(0xF2, 0xF2, 0xF4)   # 表格底色
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_THEAD  = RGBColor(0x1A, 0x1A, 0x2E)   # 表头底色

# ── 工具函数 ──
def set_cell_bg(cell, hex_color: str):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, style in kwargs.items():
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), style.get('val', 'single'))
        border.set(qn('w:sz'), str(style.get('sz', 4)))
        border.set(qn('w:color'), style.get('color', 'auto'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(text, level=1, color=None, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = C_DARK
        # 下划线装饰
        p.paragraph_format.border_bottom = None
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = C_RED if color is None else color
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = C_DARK
    return p

def body(text, indent=False, color=None, bold_parts=None):
    """普通段落，支持简单加粗标记 **text**"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    if bold_parts is None:
        # 全文解析 **…** 加粗
        import re
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                if color: run.font.color.rgb = color
            else:
                run = p.add_run(part)
                if color: run.font.color.rgb = color
        for run in p.runs:
            run.font.size = Pt(10.5)
    return p

def bullet(items, symbol='•'):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent   = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after   = Pt(3)
        p.paragraph_format.space_before  = Pt(0)
        import re
        parts = re.split(r'(\*\*[^*]+\*\*)', f'{symbol}  {item}')
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('OpenClaw 类产品商业化策略研究报告')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = C_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('C 端陷阱与 B 端机会全景分析')
run.font.size = Pt(13)
run.font.color.rgb = C_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run('2026 年 4 月  |  内部研究文件')
run.font.size = Pt(10)
run.font.color.rgb = C_GRAY

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据来源：TrustMRR · PANews · 爱建证券研究报告 · 光锥智能 · DoNews')
run.font.size = Pt(9)
run.font.color.rgb = C_GRAY
run.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════
# 第一章：背景与产品定位
# ════════════════════════════════════════════════
heading('一、背景与产品定位', 1)

heading('1.1  OpenClaw 是什么', 2)
body('OpenClaw（原名 Clawdbot / Moltbot）由 PSPDFKit 创始人 Peter Steinberger 于 2025 年 11 月启动，定义为"具备主动执行能力的 AI Agent"。项目的核心设计目标是为大语言模型赋予自主决策与任务执行能力，依托开源社区生态实现产品与用户的协同迭代。')
body('**关键时间节点：**')
bullet([
    '2024 年 4 月：Peter Steinberger 首次构思，受限于当时模型能力暂时搁置',
    '2025 年 11 月：重启研发，以 Clawdbot 为名开源，初步渗透开发者社区',
    '2026 年 1 月 28 日：腾讯云、阿里云同步上线一键部署方案',
    '2026 年 1 月 30 日：正式定名为 OpenClaw',
    '2026 年 2 月 2 日：GitHub 星标数达到 142,856，两周内突破 17 万',
])

heading('1.2  产品属性的本质判断', 2)
body('**OpenClaw 的本质是基础设施层，而非开箱即用的消费品。** 这一属性决定了其 C 端变现的天花板，也预示着 B 端定制服务的更大空间。')
body('区别于 ChatGPT / Claude 仅对话交互，OpenClaw 支持：系统级任务执行、完全本地化数据存储、10 + 聊天应用多渠道接入、主动提醒与开源可扩展能力。')

divider()

# ════════════════════════════════════════════════
# 第二章：C 端市场洞察
# ════════════════════════════════════════════════
heading('二、C 端市场洞察', 1)

heading('2.1  当前 C 端已跑通的变现场景', 2)
body('基于 TrustMRR 平台数据，OpenClaw 生态 153 个项目近 30 天合计营收约 **358,600 美元**，但 **Top 30 占比 97.3%**，高度集中于少数头部玩家。已验证的 C 端变现路径如下：')

# 表格
tbl = doc.add_table(rows=6, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ['变现类型', '代表案例', '收入规模', '可持续性']
for i, h in enumerate(hdr):
    cell = tbl.cell(0, i)
    set_cell_bg(cell, '1A1A2E')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = C_WHITE
    run.font.size = Pt(10)
rows_data = [
    ['托管部署/一键云端化', 'QuickClaw（$3.99/周）', '$120,100/月·占34.5%', '⚠ 窗口期业务，大厂入局后萎缩'],
    ['Skill包/工作流模板', 'ClawMart（Felix创建）', '累计$71,300', '⚠ 需私有知识嵌入方可持续'],
    ['个人Agent品牌化', 'FelixCraft（Nat Eliason）', '1周$3,500 Stripe', '⚠ 本质是流量叙事，难以复制'],
    ['内容自动化降本', '"Larry"（Oliver Henry）', '5天50万播放+$588', '⚠ 故事本身比产品更值钱'],
    ['硬件+深度定制', 'RoofClaw', '30天$49,800·累计$1.8M', '✓ 最健康的商业模式'],
]
for r_idx, row_data in enumerate(rows_data):
    for c_idx, val in enumerate(row_data):
        cell = tbl.cell(r_idx+1, c_idx)
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if r_idx % 2 == 0:
            set_cell_bg(cell, 'F8F8FA')

doc.add_paragraph()

heading('2.2  C 端三大结构性阻碍', 2)

body('**① 产品属性错配：基础设施不适合直接做消费品**')
body('OpenClaw 对非技术用户门槛极高，用户需要自行配置 LLM API、Skill、工作流。这决定了 C 端能买单的场景，本质上是在为"信息差"付费，而非为产品本身的持续价值付费。', indent=True)

body('**② 大厂全面入局，认知护城河快速消失**')
bullet([
    '2026 年 1 月 28 日：腾讯云/阿里云推出一键部署方案',
    '2026 年 1 月：月之暗面推出 Kimi Claw（限 199 元/月以上订阅用户）',
    '2026 年 2 月：MiniMax 推出 MaxClaw',
    '大厂方案上线后，代装/托管类 C 端服务窗口期将快速关闭',
])

body('**③ Skill 天然趋向开源，付费内容壁垒脆弱**')
body('ClawHub 的开源分发逻辑决定了优质 Skill 一经流传，复制成本趋近于零。付费 Skill 长期面临免费版挤压，类比 iOS App Store 里付费应用被免费版取代的规律。', indent=True)

body('**核心结论（PANews 调研）：** 赚最多钱的不是用 OpenClaw 做产品的人，而是帮别人部署、教别人使用、炒话题流量的人。"讲故事"本身是当前最稳定的 C 端变现路径，这不是健康的产品商业模式。')

divider()

# ════════════════════════════════════════════════
# 第三章：四大可触达人群机会评估
# ════════════════════════════════════════════════
heading('三、四大可触达人群机会评估', 1)
body('评估维度：付费意愿 × 需求强度 × Skill 可复制性 × 竞争烈度 × 决策链条长短')

tbl2 = doc.add_table(rows=5, cols=5)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = ['行业人群', '评级', '核心机会', '主要障碍', '建议切入']
for i, h in enumerate(hdr2):
    cell = tbl2.cell(0, i)
    set_cell_bg(cell, '1A1A2E')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = C_WHITE
    run.font.size = Pt(9.5)
rows2 = [
    ['① 金融从业者\n（最优先）', '★★★★★',
     '财报摘要、合规文件、KYC 整理；个人付费能力最强，AI 接受度高',
     '大机构数据安全要求高，需本地部署；采购周期 2-6 个月',
     '先做私募/独立理财师 ¥299-599/月订阅，积累案例后撬动中小机构'],
    ['② 电商白领\n（次优先）', '★★★★',
     '竞品监控、批量文案改写、选品报告；工具接受度高，需求标准化',
     '大厂 IT 管控；个人预算上限 ¥199/月；Coze 免费模板竞争',
     '个人效率订阅 ¥99-199/月，主打"比 Coze 模板更精准"'],
    ['③ 食品制造加工', '★★★',
     '食品标签合规检查（对照 GB 标准）、供应商资质管理；信息差红利大',
     '决策链长（老板/主管买单）；工厂 IT 设施弱；需行业 Know-how',
     '单次定制交付 ¥5,000-15,000，以合规风险为切入点'],
    ['④ 餐饮行业\n（慎重）', '★★',
     '连锁品牌总部有运营手册/选址分析需求（50 家以上门店）',
     '净利率 5-8%，付费极敏感；美团/商米等 SaaS 已覆盖基础需求',
     '无深度餐饮人脉则不建议作为突破口'],
]
for r_idx, row_data in enumerate(rows2):
    for c_idx, val in enumerate(row_data):
        cell = tbl2.cell(r_idx+1, c_idx)
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if r_idx % 2 == 0:
            set_cell_bg(cell, 'F8F8FA')
        if c_idx == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
divider()

# ════════════════════════════════════════════════
# 第四章：Skill 市场化路径
# ════════════════════════════════════════════════
heading('四、Skill 市场化：护城河分析与变现路径', 1)

heading('4.1  Skill 付费护城河判断', 2)

tbl3 = doc.add_table(rows=6, cols=3)
tbl3.style = 'Table Grid'
for i, h in enumerate(['护城河类型', '强度', '判断依据']):
    cell = tbl3.cell(0, i)
    set_cell_bg(cell, '1A1A2E')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = C_WHITE
    run.font.size = Pt(10)
rows3 = [
    ['信息差（我不知道这个 Skill）', '❌ 极弱', '开源后立刻消失，持续时间 1-4 周'],
    ['配置复杂（需专业调试）', '⚠ 中等', '随官方一键工具完善而下降'],
    ['私有行业数据/Prompt 嵌入', '✓ 强', '竞争者无法复制，是核心壁垒'],
    ['持续维护+模型更新同步', '✓ 强', '订阅逻辑的根基，黏性强'],
    ['行业专家品牌信任', '✓ 强', '需要时间积累，一旦建立难以被取代'],
]
for r_idx, row_data in enumerate(rows3):
    for c_idx, val in enumerate(row_data):
        cell = tbl3.cell(r_idx+1, c_idx)
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if r_idx % 2 == 0:
            set_cell_bg(cell, 'F8F8FA')

doc.add_paragraph()

heading('4.2  三条 Skill 变现路径', 2)
body('**路径一：垂直行业 Skill 包（推荐首选）**')
bullet([
    '将某行业最高频、最费时的 3-5 个任务封装为 Skill，一次开发多次销售',
    '定价：Gumroad / 爱发电，$29-99 一次性购买，或 ¥19/月订阅',
    '核心验证标准：这个 Skill 能替代多少小时人工？低于 2 小时则付费意愿极弱',
])
body('**路径二：Skill 订阅服务（中期稳现金流）**')
bullet([
    '每月新增 2-3 个经过验证的行业 Skill，定价 ¥99-299/月',
    '配合专属社群，用户可提需求，形成口碑传播',
    '需要至少 3 个月内容储备才能启动，适合有持续输出能力的团队',
])
body('**路径三：企业专属 Skill 定制（快速变现但难规模化）**')
bullet([
    '单次交付定价 ¥5,000-30,000，包含需求调研、开发调试、部署培训',
    '作为冷启动阶段的现金流来源，同时积累行业案例',
    '缺点：边际成本不降低，需配合路径一/二才有规模化潜力',
])

divider()

# ════════════════════════════════════════════════
# 第五章：两种商业模式及落地路径
# ════════════════════════════════════════════════
heading('五、两种核心商业模式及落地路径', 1)

heading('5.1  模式 A：白领数字分身', 2)
body('**产品定义：** 不是"帮你用 OpenClaw"，而是把用户的工作方式复刻进 Agent——思维框架、Prompt 偏好、输出模板——打包成专属工作流。')
body('**目标用户：** 电商运营、内容运营、产品经理等大厂白领（决策快、接受度高、有个人付费习惯）')
body('**完整交付物包含 4 层：**')
bullet([
    '**输入层**：10-20 个高频场景 Prompt 模板库',
    '**记忆层**：个人知识库搭建（公司背景、历史决策、个人偏好）',
    '**输出层**：符合个人/公司风格的输出格式规范',
    '**工作流层**：3-5 个最高频任务的自动化流程',
])
body('**收入结构：**')
bullet([
    '定制单价：¥999-3,000 / 人',
    '课程定价（沉淀后）：¥199-499 一次性购买',
    '订阅制（长期）：¥99-199/月',
])
body('**落地步骤：**')
bullet([
    '**Step 1**：先做定制（5-10 单），再把定制经验提炼成课程',
    '**Step 2**：找 3 个免费体验用户，做出完整数字分身案例，收集反馈与截图',
    '**Step 3**：用案例在小红书/朋友圈发布，定价从 ¥999 起步（定价影响认知）',
    '**Step 4**：积累 20+ 不同职业案例后，录制课程上架小报童/知识星球',
])

heading('5.2  模式 B：特定行业深度工作流定制 + 托管', 2)
body('**产品定义：** 不卖工具，卖"已经跑通的 AI 工作员"——帮企业把特定业务流程自动化，部署好、调试好、维护好，收月租。')
body('**优先目标行业：** 金融机构中小型（私募/资产管理）→ 食品制造合规方向')
body('**定价结构：**')

tbl4 = doc.add_table(rows=4, cols=3)
tbl4.style = 'Table Grid'
for i, h in enumerate(['阶段', '内容', '收费']):
    cell = tbl4.cell(0, i)
    set_cell_bg(cell, '1A1A2E')
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = C_WHITE
    run.font.size = Pt(10)
rows4 = [
    ['立项阶段', '业务流程调研（2-3 次）+ 识别前 3 个可自动化工作流 + 交付方案文档', '¥3,000-8,000（一次性）'],
    ['开发阶段', '工作流开发 + 测试 + 部署 + 核心用户培训（1-2 人）', '含在年费/托管费中'],
    ['托管阶段', '每月维护/迭代 2 次 + 模型更新同步 + 月度使用报告', '¥2,000-5,000/月'],
]
for r_idx, row_data in enumerate(rows4):
    for c_idx, val in enumerate(row_data):
        cell = tbl4.cell(r_idx+1, c_idx)
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if r_idx % 2 == 0:
            set_cell_bg(cell, 'F8F8FA')

doc.add_paragraph()
body('**落地步骤：**')
bullet([
    '**Step 1**：选定第一个标杆客户（中小型私募基金/食品公司老板，能直接拍板）',
    '**Step 2**：谈法不说"帮你部署 OpenClaw"，而说"帮你把 XX 流程自动化，省 X 人×X 小时/周，收月租"',
    '**Step 3**：交付后整理案例文章（匿名数字化），发知乎/36 氪，作为 B 端获客素材',
    '**Step 4**：做完 3-5 个同行业客户后，提炼"行业基础包"（70% 通用 + 30% 定制），降低边际成本',
])
body('**参考标杆：RoofClaw**——MacBook 预装 + Agent 定制训练服务，累计营收已达 **$1.8M**（约 1,300 万人民币）。')

heading('5.3  两种模式的协同关系', 2)
body('两种模式并非独立运作，存在明确的互相导流与背书关系：')
bullet([
    '**模式 A → 模式 B**：电商白领中若有金融/食品行业背景者，可直接转介绍到模式 B',
    '**模式 B → 模式 A**：企业客户的个人从业者（理财经理/食品工程师），可引导购买模式 A 课程',
    '**案例共用**：两条线的交付案例均可转化为内容素材，形成"内容 → 信任 → 转化"飞轮',
])

divider()

# ════════════════════════════════════════════════
# 第六章：行动建议
# ════════════════════════════════════════════════
heading('六、立即执行的行动建议', 1)

body('基于上述研究，建议团队按以下优先级推进：')

body('**行动一（本周）：找 3 个免费体验用户，做模式 A 的首批数字分身**')
bullet([
    '从电商白领中找 3 个愿意体验的人，完整交付"数字分身"，收集使用反馈和截图',
    '这是最强的冷启动信任素材，也是验证产品可行性的最低成本方式',
])

body('**行动二（两周内）：在金融人脉中找 1 个决策人做需求调研**')
bullet([
    '目标：私募研究员或独立理财顾问，摸清其最费时的重复性文档工作',
    '不急于谈钱，先做需求理解，再给方案',
])

body('**行动三（本周）：确定团队分工**')
bullet([
    '模式 A 需要内容/讲师能力；模式 B 需要销售/行业理解能力',
    '两条线不适合同一人同时推进，需明确各自负责人',
    '建议：先以模式 A 快速获得现金流，再用 1-2 个模式 B 标杆案例建立行业信用',
])

body('**战略节奏建议：**')
bullet([
    '**短期（0-3 个月）**：模式 A 电商白领订阅，快速冷启动，验证 PMF',
    '**中期（3-9 个月）**：模式 B 金融行业，做 1-2 个深度标杆案例，建立行业信用',
    '**长期（9 个月+）**：食品制造合规方向，用行业基础包降低边际成本，形成可规模化的 SaaS+服务模式',
])

divider()

# ════════════════════════════════════════════════
# 附录
# ════════════════════════════════════════════════
heading('附录：关键数据索引', 1)
body('以下数据均来自多方交叉验证，可作为内部决策参考依据：')
bullet([
    'TrustMRR OpenClaw 生态 153 个项目，近 30 天合计营收约 $358,600',
    'Top 30 项目合计营收占总体 97.3%，高度头部集中',
    'RoofClaw：近 30 天 $49,800，**累计 $1.8M**（硬件+深度定制模式）',
    'ClawMart（Skill 市场）：累计 $71,300，TrustMRR OpenClaw 类目第一',
    'QuickClaw（托管部署）：近 30 天 $8,782（$3.99/周）',
    '托管/部署类服务：占生态月收入 $120,100（34.5%）',
    'Kimi K2.5 在 OpenRouter Tokens 消耗环比增长 23,357%，OpenClaw 是主要调用来源',
    '月之暗面 C 端订阅年收入估算约 ¥2 亿',
    'GitHub 星标数（2026.02.02）：142,856，两周突破 17 万',
    '国内代装服务：闲鱼远程 ¥100-300，上门 ¥400-1,000，月环比增长 150%',
    '来源：PANews·TrustMRR / 爱建证券研究报告（2026.02.03）/ 光锥智能（2026.02）/ DoNews（2026.02.28）',
])

# 保存
out_path = '/Users/baixiaoyang/Desktop/36kr/study-career/openclaw-strategy-report.docx'
doc.save(out_path)
print(f'✅ 已生成：{out_path}')
